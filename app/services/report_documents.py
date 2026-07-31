from __future__ import annotations

import hashlib
import io
import json
import re
import uuid
from datetime import datetime
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any

from sqlalchemy import func, select
from sqlalchemy.orm import Session, joinedload

from app.config import get_settings
from app.models import (
    AnalysisJob,
    Appeal,
    AppealAnnotation,
    PolicyDocument,
    QuarterSnapshot,
    Region,
    ReportDocument,
    ReportRevision,
)
from app.services.analytics import (
    dashboard_stats,
    grouped_comparison,
    previous_quarter,
)
from app.services.privacy import redact_text
from app.services.providers import get_chat_provider
from app.schemas import ReportFactPack
from app.services.taxonomy import active_taxonomy


def _json_safe(value: object) -> object:
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, dict):
        return {str(key): _json_safe(item) for key, item in value.items()}
    if isinstance(value, (list, tuple)):
        return [_json_safe(item) for item in value]
    return value


def create_policy(
    session: Session,
    *,
    title: str,
    issuing_authority: str = "",
    source_url: str = "",
    applicable_region: str = "全国",
    published_at: datetime | None = None,
    effective_until: datetime | None = None,
    content: str = "",
    archived_path: str = "",
) -> PolicyDocument:
    normalized = "\n".join(line.strip() for line in content.splitlines() if line.strip())
    digest = hashlib.sha256(normalized.encode("utf-8")).hexdigest()
    previous_version = session.scalar(
        select(func.max(PolicyDocument.version)).where(PolicyDocument.title == title.strip())
    )
    policy = PolicyDocument(
        title=title.strip(),
        issuing_authority=issuing_authority.strip(),
        source_url=source_url.strip(),
        applicable_region=applicable_region.strip() or "全国",
        published_at=published_at,
        effective_until=effective_until,
        version=int(previous_version or 0) + 1,
        content=normalized,
        content_hash=digest,
        archived_path=archived_path,
    )
    session.add(policy)
    session.commit()
    return policy


def _representative_cases(
    session: Session,
    quarter: str,
    *,
    city: str = "",
    topics: list[str] | None = None,
    limit: int = 10,
) -> list[dict[str, object]]:
    conditions: list[object] = [
        Appeal.quarter == quarter,
        Appeal.is_canonical.is_(True),
        Appeal.reply_content.is_not(None),
        Appeal.reply_content != "",
    ]
    if city:
        conditions.append((Region.prefecture_city == city) | (Region.city == city))
    if topics:
        conditions.append(AppealAnnotation.topic.in_(topics))
    rows = session.scalars(
        select(Appeal)
        .join(Region)
        .outerjoin(AppealAnnotation)
        .options(joinedload(Appeal.annotation), joinedload(Appeal.region))
        .where(*conditions)
        .order_by(
            AppealAnnotation.confidence.desc().nullslast(),
            Appeal.received_at.desc(),
        )
        .limit(limit)
    ).unique().all()
    return [
        {
            "source_id": appeal.external_id,
            "date": appeal.received_at.date().isoformat(),
            "city": appeal.region.prefecture_city or appeal.region.city,
            "topic_l1": appeal.annotation.topic if appeal.annotation else "其他/综合",
            "topic_l2": appeal.annotation.subtopic if appeal.annotation else "",
            "title": redact_text(appeal.redacted_title)[:120],
            "content": redact_text(appeal.redacted_content)[:320],
            "department": appeal.department,
            "reply": redact_text(appeal.redacted_reply)[:320],
        }
        for appeal in rows
    ]


def _topic_share_map(stats: dict[str, object]) -> dict[str, float]:
    total = int(stats.get("event_count") or 0)
    if not total:
        return {}
    return {
        str(item["name"]): round(int(item["count"]) / total * 100, 2)
        for item in stats.get("topics", [])
    }


def _subtopic_share_map(stats: dict[str, object]) -> dict[str, float]:
    total = int(stats.get("event_count") or 0)
    if not total:
        return {}
    return {
        str(item["name"]): round(int(item["count"]) / total * 100, 2)
        for item in stats.get("subtopics", [])
        if item["name"] not in {"", "未填写"}
    }


def _consensus_suggestions(
    session: Session,
    quarter: str,
    *,
    city: str = "",
    limit: int = 15,
) -> list[dict[str, object]]:
    city_name = func.coalesce(func.nullif(Region.prefecture_city, ""), Region.city)
    conditions: list[object] = [
        Appeal.quarter == quarter,
        Appeal.appeal_type == "建议",
        Appeal.duplicate_group_key != "",
    ]
    if city:
        conditions.append(city_name == city)
    rows = session.execute(
        select(
            Appeal.duplicate_group_key,
            func.min(Appeal.redacted_title),
            func.min(AppealAnnotation.topic),
            func.min(AppealAnnotation.subtopic),
            func.count(Appeal.id),
            func.count(func.distinct(city_name)),
        )
        .join(Region)
        .outerjoin(AppealAnnotation)
        .where(*conditions)
        .group_by(Appeal.duplicate_group_key)
        .order_by(func.count(Appeal.id).desc())
        .limit(limit)
    ).all()
    return [
        {
            "group_id": str(group_id),
            "title": title or "未命名建议",
            "topic_l1": topic or "其他/综合",
            "topic_l2": subtopic or "",
            "message_count": int(message_count),
            "city_count": int(city_count),
        }
        for group_id, title, topic, subtopic, message_count, city_count in rows
    ]


def build_report_fact_pack(
    session: Session,
    *,
    report_type: str,
    quarter: str,
    city: str = "",
    city_code: str = "",
    topic_l1: str = "",
    appeal_type: str = "",
    policy_ids: list[int] | None = None,
) -> dict[str, object]:
    if report_type not in {"national", "city"}:
        raise ValueError("报告类型只能是 national 或 city")
    if report_type == "city" and not city:
        if city_code:
            city = (
                session.scalar(
                    select(Region.prefecture_city).where(Region.city_code == city_code).limit(1)
                )
                or session.scalar(select(Region.city).where(Region.city_code == city_code).limit(1))
                or ""
            )
        if not city:
            raise ValueError("城市报告必须指定 city 或有效 city_code")

    snapshot = session.scalar(
        select(QuarterSnapshot)
        .where(QuarterSnapshot.quarter == quarter, QuarterSnapshot.status == "active")
        .order_by(QuarterSnapshot.version.desc())
    )
    taxonomy = active_taxonomy(session)
    scope_stats = dashboard_stats(
        session,
        city=city or None,
        quarter=quarter,
        topic_l1=topic_l1 or None,
        appeal_type=appeal_type or None,
    )
    national_stats = dashboard_stats(
        session,
        quarter=quarter,
        topic_l1=topic_l1 or None,
        appeal_type=appeal_type or None,
    )
    previous_stats = dashboard_stats(
        session,
        city=city or None,
        quarter=previous_quarter(quarter),
        topic_l1=topic_l1 or None,
        appeal_type=appeal_type or None,
    )
    top_topics = [str(item["name"]) for item in scope_stats["topics"][:5]]
    cases = _representative_cases(
        session,
        quarter,
        city=city,
        topics=top_topics,
        limit=10 if report_type == "national" else 6,
    )
    policies = []
    if policy_ids:
        policies_statement = select(PolicyDocument).where(
            PolicyDocument.status == "active",
            PolicyDocument.id.in_(policy_ids),
        )
        policies = session.scalars(
            policies_statement.order_by(PolicyDocument.published_at.desc()).limit(20)
        ).all()
    policy_facts = [
        {
            "source_id": str(policy.id),
            "title": policy.title,
            "issuing_authority": policy.issuing_authority,
            "published_at": policy.published_at.date().isoformat() if policy.published_at else "",
            "applicable_region": policy.applicable_region,
            "source_url": policy.source_url,
            "excerpt": redact_text(policy.content)[:700],
        }
        for policy in policies
    ]

    pack: dict[str, object] = {
        "schema_version": "report-fact-pack-v1",
        "report_type": report_type,
        "quarter": quarter,
        "scope": {
            "city": city,
            "city_code": city_code,
            "topic_l1": topic_l1,
            "appeal_type": appeal_type,
            "count_basis": "deduplicated_events",
        },
        "snapshot": {
            "id": snapshot.id if snapshot else None,
            "version": snapshot.version if snapshot else None,
            "status": snapshot.status if snapshot else "live-data-not-frozen",
        },
        "taxonomy": {
            "id": taxonomy.id,
            "version": taxonomy.version,
            "status": taxonomy.status,
            "l1_macro_f1": taxonomy.l1_macro_f1,
            "l2_macro_f1": taxonomy.l2_macro_f1,
        },
        "statistics": _json_safe(scope_stats),
        "previous_quarter_statistics": _json_safe(previous_stats),
        "cases": cases,
        "policies": policy_facts,
    }
    if report_type == "national":
        pack["comparisons"] = {
            dimension: grouped_comparison(
                session, quarter, dimension, topic_l1=topic_l1 or None
            )
            for dimension in ("macro_region", "city_tier", "urban_rural")
        }
        pack["consensus_suggestions"] = _consensus_suggestions(session, quarter)
    else:
        city_shares = _subtopic_share_map(scope_stats)
        national_shares = _subtopic_share_map(national_stats)
        previous_shares = _subtopic_share_map(previous_stats)
        distinctive = [
            {
                "topic": topic,
                "city_share": share,
                "national_share": national_shares.get(topic, 0),
                "difference": round(share - national_shares.get(topic, 0), 2),
            }
            for topic, share in city_shares.items()
        ]
        distinctive.sort(key=lambda item: abs(float(item["difference"])), reverse=True)
        topic_changes = [
            {
                "topic": topic,
                "previous_share": previous_shares.get(topic, 0),
                "current_share": share,
                "change": round(share - previous_shares.get(topic, 0), 2),
            }
            for topic, share in city_shares.items()
        ]
        topic_changes.sort(key=lambda item: abs(float(item["change"])), reverse=True)
        pack["national_benchmark"] = _json_safe(national_stats)
        pack["distinctive_topics"] = distinctive[:10]
        pack["topic_changes"] = topic_changes[:4] if int(previous_stats["event_count"]) else []
    return ReportFactPack.model_validate(pack).model_dump(exclude_none=True)


def _ranking_lines(rows: list[dict[str, object]], limit: int = 10) -> str:
    return "\n".join(
        f"{index}. {item['name']}：{item['count']}件"
        for index, item in enumerate(rows[:limit], start=1)
    )


def _case_lines(cases: list[dict[str, object]], limit: int) -> str:
    return "\n".join(
        f"- [{item['source_id']}] {item['city']}，{item['topic_l1']}："
        f"{item['title']}；回复部门：{item['department']}。"
        for item in cases[:limit]
    ) or "暂无满足条件的可引用案例。"


def _policy_lines(policies: list[dict[str, object]]) -> str:
    return "\n".join(
        f"- [政策:{item['source_id']}]《{item['title']}》，{item['issuing_authority']}，"
        f"{item['published_at'] or '日期未录入'}。"
        for item in policies
    ) or "本期未绑定权威政策材料，不作具体政策引用。"


def _percentage_or_unavailable(value: object | None) -> str:
    return f"{value}%" if value is not None else "暂无"


def _quality_lines(dimensions: list[dict[str, object]]) -> str:
    return "\n".join(
        f"- {item['name']}：达成{item['yes']}件、未达成{item['no']}件、"
        f"不适用{item['not_applicable']}件、未知{item['unknown']}件；"
        f"适用样本达成率{_percentage_or_unavailable(item['yes_rate'])}。"
        for item in dimensions
    ) or "暂无回复质量分项数据。"


def _suggestion_lines(items: list[dict[str, object]]) -> str:
    return "\n".join(
        f"- {item['title']}（{item['topic_l2'] or item['topic_l1']}）："
        f"涉及{item['message_count']}条留言、{item['city_count']}个城市。"
        for item in items
    ) or "本期暂无可稳定归并的高共识建议。"


def _topic_change_lines(items: list[dict[str, object]]) -> str:
    return "\n".join(
        f"- {item['topic']}：{item['previous_share']}% → {item['current_share']}%，"
        f"变化{float(item['change']):+.2f}个百分点。"
        for item in items
    ) or "无可比上一季度，暂不展示议题变化。"


def local_report_from_fact_pack(title: str, pack: dict[str, object]) -> str:
    stats = pack["statistics"]
    previous = pack["previous_quarter_statistics"]
    current_total = int(stats["event_count"])
    previous_total = int(previous["event_count"])
    growth = (
        round((current_total - previous_total) / previous_total * 100, 1)
        if previous_total
        else None
    )
    growth_text = f"{growth:+.1f}%" if growth is not None else "无可比基期"
    report_type = str(pack["report_type"])
    shared = f"""# {title}

## 一、总体判断

本期按去重事件口径纳入{current_total}件，原始留言{stats['raw_total']}件，重复率{stats['duplicate_rate']}%。相较上一季度，事件量变化为{growth_text}。回复率为{stats['response_rate']}%，平均回复耗时为{stats['average_response_hours'] if stats['average_response_hours'] is not None else '暂无'}小时。

## 二、诉求主题结构

{_ranking_lines(stats['topics'], 10) or '暂无主题统计。'}

### 重点细分问题

{_ranking_lines(stats['subtopics'], 10) or '暂无已审核的细分问题统计。'}
"""
    if report_type == "national":
        comparisons = pack.get("comparisons", {})
        comparison_text = "\n\n".join(
            f"### {name}\n\n"
            + "\n".join(
                f"- {item['name']}：{item['event_count']}件，回复五项平均达成率"
                f"{_percentage_or_unavailable(item['reply_quality_score'])}"
                for item in rows
            )
            for name, rows in (
                ("宏观区域比较", comparisons.get("macro_region", [])),
                ("城市行政层级比较", comparisons.get("city_tier", [])),
                ("城乡比较", comparisons.get("urban_rural", [])),
            )
        )
        body = f"""
## 三、区域与城市层级比较

{comparison_text or '暂无可比较数据。'}
"""
        tail = f"""
## 四、政府回应和诉求办理情况

本期已回复{stats['responded']}件，未回复{stats['pending']}件。回复速度与回复内容质量分开观察，五项结果如下：

{_quality_lines(stats.get('reply_quality_dimensions', []))}

上述结果保留“不适用”和“未知”状态，研究判断需回到证据片段复核，不以模型单一总分替代。

## 五、典型案例

{_case_lines(pack.get('cases', []), 5)}

## 六、政策依据

{_policy_lines(pack.get('policies', []))}

## 专栏：问计于民

{_suggestion_lines(pack.get('consensus_suggestions', []))}

## 七、治理建议

建议围绕高频问题建立季度专题台账，对重复点位、重复主体和跨部门事项进行归并；将回复中的核查、措施、时限、责任与后续渠道纳入复核；对明显高于全国平均的城市特色问题形成专项调研。

## 八、数据说明

本报告使用季度冻结快照、去重事件主口径和版本化标签生成。标签状态为“{pack['taxonomy']['status']}”；未通过黄金样本验收前，结论仅作为试运行分析。
"""
    else:
        distinctive = pack.get("distinctive_topics", [])
        leading_department = (
            stats["departments"][0]["name"] if stats.get("departments") else "暂无稳定数据"
        )
        body = f"""
## 三、城市画像卡片

- 城市：{pack['scope']['city']}
- 统计时段：{pack['quarter']}
- 去重事件量：{stats['event_count']}件
- 主要承办部门：{leading_department}

## 四、本市特色问题

""" + (
            "\n".join(
                f"- {item['topic']}：本市占比{item['city_share']}%，全国占比{item['national_share']}%，相差{item['difference']:+.2f}个百分点。"
                for item in distinctive
            )
            or "暂无可比较的城市特色主题。"
        ) + f"""

## 五、本季变化

{_topic_change_lines(pack.get('topic_changes', []))}
"""
        tail = f"""
## 六、政府回应和办理质量

本期已回复{stats['responded']}件，未回复{stats['pending']}件。回复速度与回复内容质量分开观察，五项结果如下：

{_quality_lines(stats.get('reply_quality_dimensions', []))}

上述结果保留“不适用”和“未知”状态，研究判断需回到证据片段复核，不以模型单一总分替代。

## 七、典型市民声音

{_case_lines(pack.get('cases', []), 2)}

## 八、政策依据

{_policy_lines(pack.get('policies', []))}

## 九、治理建议

建议围绕高频问题建立季度专题台账，对重复点位、重复主体和跨部门事项进行归并；将回复中的核查、措施、时限、责任与后续渠道纳入复核；对明显高于全国平均的城市特色问题形成专项调研。

## 十、数据说明

本报告使用季度冻结快照、去重事件主口径和版本化标签生成。标签状态为“{pack['taxonomy']['status']}”；未通过黄金样本验收前，结论仅作为试运行分析。
"""
    return shared + body + tail


def _append_source_appendix(content: str, pack: dict[str, object]) -> str:
    appendix = "\n\n## 证据索引\n\n"
    appendix += "### 留言案例\n\n" + _case_lines(pack.get("cases", []), 10)
    appendix += "\n\n### 政策材料\n\n" + _policy_lines(pack.get("policies", []))
    return content.rstrip() + appendix + "\n"


_NUMERIC_CLAIM_RE = re.compile(r"(?<![A-Za-z0-9])[-+]?\d+(?:\.\d+)?%?")


def _normalized_number(value: str) -> str:
    try:
        return str(Decimal(value.rstrip("%").lstrip("+")).normalize())
    except InvalidOperation:
        return value


def validate_report(content: str, pack: dict[str, object]) -> list[str]:
    failures: list[str] = []
    if pack.get("snapshot", {}).get("status") != "active":
        failures.append("报告未绑定已激活的季度快照")
    stats = pack["statistics"]
    for value, label in (
        (stats["event_count"], "去重事件量"),
        (stats["raw_total"], "原始留言量"),
        (stats["response_rate"], "回复率"),
    ):
        if str(value) not in content:
            failures.append(f"正文未包含锁定事实：{label}={value}")
    for case in pack.get("cases", [])[:2]:
        if f"[{case['source_id']}]" not in content:
            failures.append(f"正文未引用案例：{case['source_id']}")
    for policy in pack.get("policies", []):
        if f"[政策:{policy['source_id']}]" not in content:
            failures.append(f"正文未引用政策：{policy['source_id']}")
    policy_ids = {str(item["source_id"]) for item in pack.get("policies", [])}
    unsupported_policy_ids = sorted(
        set(re.findall(r"\[政策:([^\]]+)\]", content)) - policy_ids
    )
    if unsupported_policy_ids:
        failures.append(f"正文引用了事实包之外的政策：{','.join(unsupported_policy_ids)}")

    allowed_numbers = {
        _normalized_number(item)
        for item in _NUMERIC_CLAIM_RE.findall(json.dumps(pack, ensure_ascii=False))
    }
    # Markdown section numbers are layout, not empirical claims.
    allowed_numbers.update(str(item) for item in range(1, 21))
    unsupported_numbers = sorted(
        {
            item
            for item in _NUMERIC_CLAIM_RE.findall(content)
            if _normalized_number(item) not in allowed_numbers
        }
    )
    if unsupported_numbers:
        failures.append(
            "正文出现事实包未支持的数字：" + "、".join(unsupported_numbers[:20])
        )
    return failures


def create_report_document(
    session: Session,
    payload: dict[str, object],
    *,
    job: AnalysisJob | None = None,
) -> ReportDocument:
    report_type = str(payload.get("report_type") or "")
    quarter = str(payload.get("quarter") or "")
    city = str(payload.get("city") or "")
    city_code = str(payload.get("city_code") or "")
    pack = build_report_fact_pack(
        session,
        report_type=report_type,
        quarter=quarter,
        city=city,
        city_code=city_code,
        topic_l1=str(payload.get("topic_l1") or ""),
        appeal_type=str(payload.get("appeal_type") or ""),
        policy_ids=[int(item) for item in payload.get("policy_ids", [])],
    )
    city = str(pack["scope"]["city"])
    title = (
        f"全国居民留言季度研究报告（{quarter}）"
        if report_type == "national"
        else f"{city}居民留言季度简报（{quarter}）"
    )
    if job:
        job.progress = 35
        session.commit()

    content = local_report_from_fact_pack(title, pack)
    generated_by = "local-template"
    provider = get_chat_provider(session=session)
    if provider.enabled and bool(payload.get("use_model", True)):
        try:
            generated = provider.complete(
                (
                    "你是社会治理研究报告撰写助手。只能依据事实包写作，不得增加事实包之外的"
                    "数字、案例、政策或因果结论。案例必须使用[来源ID]，政策必须使用"
                    "[政策:ID]。保持正式、克制、可核验的中文Markdown结构。"
                ),
                "报告标题："
                + title
                + "\n\n事实包：\n"
                + json.dumps(pack, ensure_ascii=False),
                purpose="report",
                prompt_version="report-fact-pack-v1",
            )
            candidate = _append_source_appendix(generated, pack)
            if not validate_report(candidate, pack):
                content = candidate
                generated_by = provider.model_name
            else:
                generated_by = f"local-template ({provider.model_name} validation fallback)"
        except Exception:
            generated_by = f"local-template ({provider.model_name} unavailable)"

    if "## 证据索引" not in content:
        content = _append_source_appendix(content, pack)
    report = ReportDocument(
        report_type=report_type,
        mode=str(payload.get("mode") or "standard"),
        title=title,
        quarter=quarter,
        city_code=city_code,
        city=city,
        snapshot_id=pack["snapshot"]["id"],
        taxonomy_version_id=pack["taxonomy"]["id"],
        status="draft",
        fact_pack=pack,
        current_content=content,
        generated_by=generated_by,
    )
    session.add(report)
    session.flush()
    report.revisions.append(
        ReportRevision(
            version=1,
            content=content,
            change_note="系统生成初稿",
        )
    )
    session.commit()
    return report


def pregenerate_standard_reports(
    session: Session,
    snapshot: QuarterSnapshot,
    *,
    job: AnalysisJob | None = None,
) -> dict[str, int]:
    from app.models import CityQuarterAggregate

    scopes: list[dict[str, object]] = [
        {
            "report_type": "national",
            "mode": "standard",
            "quarter": snapshot.quarter,
            "use_model": False,
        }
    ]
    cities = session.execute(
        select(
            CityQuarterAggregate.city_code,
            CityQuarterAggregate.city,
        )
        .where(
            CityQuarterAggregate.snapshot_id == snapshot.id,
            CityQuarterAggregate.topic_l1 == "",
            CityQuarterAggregate.appeal_type == "",
            CityQuarterAggregate.canonical_count > 0,
        )
        .order_by(CityQuarterAggregate.city)
    ).all()
    scopes.extend(
        {
            "report_type": "city",
            "mode": "standard",
            "quarter": snapshot.quarter,
            "city_code": city_code,
            "city": city,
            "use_model": False,
        }
        for city_code, city in cities
    )
    created = 0
    skipped = 0
    for index, payload in enumerate(scopes, start=1):
        report_type = str(payload["report_type"])
        city_code = str(payload.get("city_code") or "")
        city = str(payload.get("city") or "")
        existing = session.scalar(
            select(ReportDocument.id).where(
                ReportDocument.snapshot_id == snapshot.id,
                ReportDocument.report_type == report_type,
                ReportDocument.mode == "standard",
                ReportDocument.city_code == city_code,
                ReportDocument.city == city,
            )
        )
        if existing:
            skipped += 1
        else:
            create_report_document(session, payload)
            created += 1
        if job:
            job.progress = min(99, 90 + int(index / max(len(scopes), 1) * 9))
            job.message = f"正在预生成标准报告：{index}/{len(scopes)}"
            session.commit()
    return {"created": created, "skipped": skipped, "total": len(scopes)}


def update_report(
    session: Session,
    report: ReportDocument,
    content: str,
    *,
    change_note: str = "",
    editor_id: int | None = None,
) -> ReportDocument:
    if report.status == "published":
        raise ValueError("已发布报告不可直接修改，请基于其创建新草稿")
    next_version = (
        session.scalar(
            select(func.max(ReportRevision.version)).where(ReportRevision.report_id == report.id)
        )
        or 0
    ) + 1
    report.current_content = content
    report.revisions.append(
        ReportRevision(
            version=next_version,
            content=content,
            change_note=change_note,
            editor_id=editor_id,
        )
    )
    session.commit()
    return report


def publish_report(
    session: Session,
    report: ReportDocument,
    *,
    publisher_id: int | None = None,
) -> ReportDocument:
    failures = validate_report(report.current_content, report.fact_pack)
    if failures:
        raise ValueError("；".join(failures))
    report.status = "published"
    report.published_by = publisher_id
    report.published_at = datetime.now()
    session.commit()
    return report


_MARKDOWN_PREFIX_RE = re.compile(r"^(#{1,6}|[-*]|\d+\.)\s*")


def _plain_lines(markdown_text: str) -> list[str]:
    return [
        _MARKDOWN_PREFIX_RE.sub("", line).replace("**", "").replace("`", "").strip()
        for line in markdown_text.splitlines()
    ]


def _chart_font(size: int, *, bold: bool = False):
    from PIL import ImageFont

    candidates = (
        (
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"
            if bold
            else "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
        ),
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc",
    )
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def _topic_chart_image(report: ReportDocument) -> Path | None:
    topics = list((report.fact_pack or {}).get("statistics", {}).get("topics", []))[:10]
    if not topics:
        return None
    from PIL import Image, ImageDraw

    settings = get_settings()
    destination = settings.export_dir / f"{report.public_id}-topics.png"
    width, height = 1500, 220 + len(topics) * 82
    image = Image.new("RGB", (width, height), "#F8FAF8")
    draw = ImageDraw.Draw(image)
    title_font = _chart_font(42, bold=True)
    label_font = _chart_font(27)
    value_font = _chart_font(26, bold=True)
    muted_font = _chart_font(22)
    draw.text((72, 48), "一级主题结构（去重事件）", fill="#173D3A", font=title_font)
    draw.text(
        (72, 108),
        f"{report.quarter} · 数据来自锁定的 ReportFactPack",
        fill="#667572",
        font=muted_font,
    )
    maximum = max(int(item.get("count") or 0) for item in topics) or 1
    bar_left, bar_width = 470, 820
    for index, item in enumerate(topics):
        y = 180 + index * 82
        name = str(item.get("name") or "未填写")
        count = int(item.get("count") or 0)
        if len(name) > 16:
            name = name[:15] + "…"
        draw.text((72, y + 8), name, fill="#243B39", font=label_font)
        draw.rounded_rectangle(
            (bar_left, y + 8, bar_left + bar_width, y + 46),
            radius=14,
            fill="#E3EAE6",
        )
        fill_width = max(12, round(bar_width * count / maximum))
        draw.rounded_rectangle(
            (bar_left, y + 8, bar_left + fill_width, y + 46),
            radius=14,
            fill="#247B67",
        )
        draw.text((1320, y + 5), f"{count:,}", fill="#173D3A", font=value_font)
    image.save(destination, format="PNG", optimize=True)
    return destination


def _set_docx_run_font(run, name: str, size: float, *, bold: bool = False, color=None):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = name
    run_properties = run._element.get_or_add_rPr()
    run_properties.rFonts.set(qn("w:eastAsia"), name)
    run_properties.rFonts.set(qn("w:ascii"), name)
    run_properties.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _docx_cjk_font() -> str:
    if Path("/System/Library/Fonts/Hiragino Sans GB.ttc").exists():
        return "Hiragino Sans GB"
    if Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc").exists():
        return "Noto Sans CJK SC"
    return "Arial Unicode MS"


def _configure_docx_table(table, widths: list[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    table_properties = table._tbl.tblPr
    width = table_properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        table_properties.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(sum(widths)))
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    table_properties.append(indent)
    margins = OxmlElement("w:tblCellMar")
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    table_properties.append(margins)
    grid_columns = list(table._tbl.tblGrid.gridCol_lst)
    for grid_column, cell_width in zip(grid_columns, widths):
        grid_column.set(qn("w:w"), str(cell_width))
    for row in table.rows:
        for cell, cell_width in zip(row.cells, widths):
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(cell_width))


def _shade_docx_cells(cells, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for cell in cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shading)


def _add_docx_fact_tables(document, report: ReportDocument) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    stats = report.fact_pack["statistics"]
    overview = document.add_table(rows=2, cols=4)
    overview.style = "Table Grid"
    values = (
        ("去重事件", f"{stats['event_count']:,}", "原始留言", f"{stats['raw_total']:,}"),
        (
            "回复率",
            f"{stats['response_rate']}%",
            "平均回复耗时",
            f"{stats['average_response_hours'] or '暂无'}小时",
        ),
    )
    for row, values_row in zip(overview.rows, values):
        for cell, value in zip(row.cells, values_row):
            cell.text = str(value)
    _configure_docx_table(overview, [1400, 3280, 1400, 3280])
    _shade_docx_cells(
        [overview.cell(0, 0), overview.cell(0, 2), overview.cell(1, 0), overview.cell(1, 2)],
        "E8EEF5",
    )
    for row in overview.rows:
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = 1
            cell.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if index % 2 else WD_ALIGN_PARAGRAPH.LEFT
            )

    topics = list(stats.get("topics", []))[:10]
    if topics:
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for cell, value in zip(table.rows[0].cells, ("一级主题", "事件量", "占比")):
            cell.text = value
        _shade_docx_cells(table.rows[0].cells, "E8EEF5")
        total = int(stats["event_count"] or 0)
        for item in topics:
            cells = table.add_row().cells
            count = int(item["count"])
            values_row = (
                item["name"],
                f"{count:,}",
                f"{count / total * 100:.2f}%" if total else "0%",
            )
            for cell, value in zip(cells, values_row):
                cell.text = str(value)
        _configure_docx_table(table, [4960, 2200, 2200])


def _add_docx_quality_table(document, report: ReportDocument) -> None:
    dimensions = list(
        report.fact_pack.get("statistics", {}).get("reply_quality_dimensions", [])
    )
    if not dimensions:
        return
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headings = ("指标", "达成", "未达成", "不适用", "未知", "达成率")
    for cell, value in zip(table.rows[0].cells, headings):
        cell.text = value
    _shade_docx_cells(table.rows[0].cells, "E8EEF5")
    for item in dimensions:
        cells = table.add_row().cells
        values = (
            item["name"],
            item["yes"],
            item["no"],
            item["not_applicable"],
            item["unknown"],
            f"{item['yes_rate']}%" if item["yes_rate"] is not None else "暂无",
        )
        for cell, value in zip(cells, values):
            cell.text = str(value)
    _configure_docx_table(table, [2560, 1300, 1300, 1300, 1300, 1600])


def export_report_docx(report: ReportDocument) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt, RGBColor

    settings = get_settings()
    destination = settings.export_dir / f"{report.public_id}.docx"
    document = Document()
    document.core_properties.title = report.title
    document.core_properties.subject = "居民政策诉求季度研究报告"
    document.core_properties.author = "居民留言分析平台"
    # Named override to the standard_business_brief preset: Chinese research
    # reports use A4 and SimSun while keeping the preset's spacing and geometry.
    section = document.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)
    section.top_margin = section.right_margin = Mm(25.4)
    section.bottom_margin = section.left_margin = Mm(25.4)
    section.header_distance = section.footer_distance = Mm(12.5)
    styles = document.styles
    cjk_font = _docx_cjk_font()
    for style_name, size, color, before, after in (
        ("Normal", 11, "202D2C", 0, 6),
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
        ("List Bullet", 11, "202D2C", 0, 8),
    ):
        style = styles[style_name]
        style.font.name = cjk_font
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cjk_font)
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), cjk_font)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), cjk_font)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1 if style_name != "List Bullet" else 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run(f"居民政策诉求季度研究 · {report.quarter}")
    _set_docx_run_font(header_run, cjk_font, 9, color=RGBColor(102, 117, 114))
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("第 ")
    _set_docx_run_font(footer_run, cjk_font, 9, color=RGBColor(102, 117, 114))
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    trailing = footer.add_run(" 页")
    _set_docx_run_font(trailing, cjk_font, 9, color=RGBColor(102, 117, 114))

    chart_path = _topic_chart_image(report)
    fact_tables_added = False
    quality_table_added = False
    for raw in report.current_content.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            title = document.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.paragraph_format.space_before = Pt(28)
            title.paragraph_format.space_after = Pt(8)
            run = title.add_run(line[2:])
            _set_docx_run_font(run, cjk_font, 26, bold=True, color=RGBColor(23, 61, 58))
            metadata = document.add_paragraph()
            metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
            metadata.paragraph_format.space_after = Pt(20)
            metadata_run = metadata.add_run(
                f"{report.quarter} · {report.report_type} · {report.status}"
            )
            _set_docx_run_font(
                metadata_run,
                cjk_font,
                10,
                color=RGBColor(102, 117, 114),
            )
            _add_docx_fact_tables(document, report)
            fact_tables_added = True
            if chart_path:
                caption = document.add_paragraph("图1  一级主题结构（去重事件）")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_before = Pt(10)
                caption.paragraph_format.space_after = Pt(4)
                document.add_picture(str(chart_path), width=Inches(6.25))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            heading = document.add_heading(line[3:], level=1)
            if "政府回应" in line and not quality_table_added:
                _add_docx_quality_table(document, report)
                quality_table_added = True
        elif line.startswith("### "):
            document.add_heading(line[4:], level=2)
        elif line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(_MARKDOWN_PREFIX_RE.sub("", line))
    if not fact_tables_added:
        _add_docx_fact_tables(document, report)
    document.save(destination)
    return destination


def export_report_pdf(report: ReportDocument) -> Path:
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import (
        Image,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    settings = get_settings()
    destination = settings.export_dir / f"{report.public_id}.pdf"
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "ChineseBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=10.5,
        leading=18,
        spaceAfter=8,
    )
    heading = ParagraphStyle(
        "ChineseHeading",
        parent=body,
        fontSize=15,
        leading=22,
        spaceBefore=12,
        spaceAfter=8,
    )
    title_style = ParagraphStyle(
        "ChineseTitle",
        parent=heading,
        fontSize=20,
        leading=28,
        alignment=TA_CENTER,
        spaceAfter=20,
    )
    caption = ParagraphStyle(
        "ChineseCaption",
        parent=body,
        fontSize=9,
        textColor=colors.HexColor("#667572"),
        alignment=TA_CENTER,
        spaceBefore=8,
        spaceAfter=5,
    )
    story: list[Any] = []
    chart_path = _topic_chart_image(report)
    quality_table_added = False
    for raw in report.current_content.splitlines():
        line = raw.strip()
        if not line:
            story.append(Spacer(1, 4))
            continue
        escaped = (
            _MARKDOWN_PREFIX_RE.sub("", line)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        if line.startswith("# "):
            story.append(Paragraph(escaped, title_style))
            stats = report.fact_pack["statistics"]
            overview = Table(
                [
                    ["去重事件", f"{stats['event_count']:,}", "原始留言", f"{stats['raw_total']:,}"],
                    [
                        "回复率",
                        f"{stats['response_rate']}%",
                        "平均回复耗时",
                        f"{stats['average_response_hours'] or '暂无'}小时",
                    ],
                ],
                colWidths=[25 * mm, 52 * mm, 28 * mm, 52 * mm],
            )
            overview.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8EEF5")),
                        ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#E8EEF5")),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CAD3CE")),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 7),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.extend([overview, Spacer(1, 9)])
            if chart_path:
                story.append(Paragraph("图1  一级主题结构（去重事件）", caption))
                from PIL import Image as PillowImage

                with PillowImage.open(chart_path) as chart_source:
                    chart_height = min(
                        115 * mm,
                        157 * mm * chart_source.height / chart_source.width,
                    )
                chart = Image(str(chart_path), width=157 * mm, height=chart_height)
                chart.hAlign = "CENTER"
                story.append(chart)
        elif line.startswith(("## ", "### ")):
            story.append(Paragraph(escaped, heading))
            if "政府回应" in line and not quality_table_added:
                dimensions = report.fact_pack["statistics"].get(
                    "reply_quality_dimensions", []
                )
                if dimensions:
                    values = [["指标", "达成", "未达成", "不适用", "未知", "达成率"]]
                    values.extend(
                        [
                            item["name"],
                            item["yes"],
                            item["no"],
                            item["not_applicable"],
                            item["unknown"],
                            f"{item['yes_rate']}%" if item["yes_rate"] is not None else "暂无",
                        ]
                        for item in dimensions
                    )
                    quality_table = Table(
                        values,
                        colWidths=[40 * mm, 21 * mm, 21 * mm, 21 * mm, 21 * mm, 29 * mm],
                        repeatRows=1,
                    )
                    quality_table.setStyle(
                        TableStyle(
                            [
                                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CAD3CE")),
                                ("ALIGN", (1, 0), (-1, -1), "CENTER"),
                                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                                ("TOPPADDING", (0, 0), (-1, -1), 5),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                            ]
                        )
                    )
                    story.extend([quality_table, Spacer(1, 8)])
                quality_table_added = True
        else:
            story.append(Paragraph(escaped, body))
    document = SimpleDocTemplate(
        str(destination),
        pagesize=A4,
        rightMargin=48,
        leftMargin=48,
        topMargin=52,
        bottomMargin=52,
        title=report.title,
    )

    def draw_page(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont("STSong-Light", 8)
        canvas.setFillColor(colors.HexColor("#667572"))
        canvas.drawString(48, A4[1] - 30, f"居民政策诉求季度研究 · {report.quarter}")
        canvas.drawCentredString(A4[0] / 2, 28, f"第 {doc.page} 页")
        canvas.restoreState()

    document.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    return destination
